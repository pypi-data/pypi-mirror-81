#!/bin/env python3
"""Module for laboratory reporting.

Containers data and interpretations together and then
streamlines it to different output formats.
"""
import matplotlib.pyplot as plt, re
from matplotlib.figure import Figure
import pandas as pd, re, inspect, os
from itertools import count
from unittest.mock import Mock

# Leopard utilities
from leopard.utils import makeFigFromFile, pdSeriesToFrame, renewliner
from leopard.utils import LeopardDict as OrderedDict
#from collections import OrderedDict

# Settings
from .config import config
reportsDir = config['leopard']['reportdir']
csvsep = config['leopard']['csvsep']
csvdec = config['leopard']['csvdec']

class Section(object):
    """
    Report sections.
    Defines methods dealing with the structure of
    sections, and section specific output.

    tablehead => set to integer to only show DataFrame.head(tablehead) rows
    in this section

    For adding subsections, a method is provided.
    """
    def __init__(self, title, text='', figures=None, tables=None,
                 subsections=None, code=None, tablehead=None,
                 tablecolumns=None, clearpage=False):
        """
        Args:
            title (str | (str,str)): Section title, or (title, subtitle).
            text (str): Section paragraph text, can also be a list of strings.
            figures (OrderedDict): Section figures.
            tables (OrderedDict): Section tables.
            subsections (list): Section subsections list.
            code (str): code that generated section data/figures, or anything to display verbatim.
            tablehead (int): Number of lines to print of tables in space restricted output formats.
            tablecolumns (list): Column names of tables included in space restricted output formats.
            clearpage (bool): Start a new page in certain output formats.
        """
        if isinstance(title, str):
            self.title, self.subtitle = title.strip(), ''
        else: self.title, self.subtitle = title[0].strip(), title[1].strip()
        self.p = text
        self.figs = OrderedDict(figures) if figures else OrderedDict()
        self.tabs = OrderedDict(tables) if tables else OrderedDict()
        self.subs = subsections if subsections else []
        self.code = code
        self.settings = {'tablehead':tablehead,
                         'tablecolumns':tablecolumns,
                         'clearpage':clearpage,
                         'doubleslashnewline':False}
        # Initially set section upstream references to self
        self._parentSection = self
        self._reportSection = self

    def __repr__(self):
        return "<{} @ {}{}>".format(
            'Report' if 'sections' in dir(self) else 'Section',
            self.title[:50],
            '' if len(self.title)<=50 else '...'
        )
    
    def __getitem__(self,key):
        try: return self.subs[key]
        except TypeError:
            try: return self._subs[key]
            except (AttributeError,KeyError) as e:
                self._subs = {s.title:s for s in self.subs}
                return self._subs[key]

    @property
    def p(self):
        return self.paragraphs

    @p.setter
    def p(self, value):
        self.paragraphs = [
            v.strip() for v in value
        ] if isinstance(value, list) else value.strip()  
    
    def append(self,*args,toSection=None,**kwargs):
        """Append a new section
        If toSection is None, section is appended to the main section/subs list.
        Else if toSection is int or (int,int,...), it gets added to the subs (subsection)
        list of the specified section.

        *args* and *kwargs* are processed by Section class initiation
        """
        if not toSection and toSection is not 0:
            s = Section(*args,**kwargs)
            self.subs.append(s)
            self.lastSection = s
            s._parentSection = self
            s._reportSection = self._reportSection
        else:
            if type(toSection) is int: toSection = (toSection,)
            s = self.subs[toSection[0]].append(*args,toSection=toSection[1:],**kwargs)
        return s

    @staticmethod
    def sectionWalker(section,callback,*args,walkTrace=tuple(),**kwargs):
        """
        callback needs to be a function that handles different 
        Section elements appropriately
        walkTrace needs to be a tuple, indicate the route to the section, e.g. (1,2,0)
        """
        callback(section,*args,walkTrace=walkTrace,case='sectionmain',**kwargs)
        c = count(1)
        for f in section.figs.items():
            callback(section,*args,walkTrace=walkTrace,case='figure',element=f,**kwargs)
        c = count(1)
        for t in section.tabs.items():
            callback(section,*args,walkTrace=walkTrace,case='table',element=t,**kwargs)
        c = count(1)
        for s in section.subs:
            Section.sectionWalker(s,callback,*args,walkTrace=walkTrace+(next(c),),**kwargs)

    def walkerWrapper(callback):
        """
        Wraps a callback function in a wrapper that will be applied to all [sub]sections.

        Returns:
            function
        """
        def wrapper(*args,**kwargs):
            #args[0] => has to be the current walked section
            return Section.sectionWalker(args[0],callback,*args[1:],**kwargs)
        return wrapper

    @walkerWrapper
    def list(self,walkTrace=tuple(),case=None,element=None):
        """List section titles.
        """
        if case == 'sectionmain': print(walkTrace,self.title)

    @walkerWrapper
    def listFigures(self,walkTrace=tuple(),case=None,element=None):
        """List section figures.
        """
        if case == 'sectionmain': print(walkTrace,self.title)
        if case == 'figure':
            caption,fig = element
            try:
                print(walkTrace,fig._leopardref,caption)
            except AttributeError:
                fig._leopardref = next(self._reportSection._fignr)
                print(walkTrace,fig._leopardref,caption)

    @walkerWrapper
    def listTables(self,walkTrace=tuple(),case=None,element=None):
        """List section tables.
        """
        if case == 'sectionmain': print(walkTrace,self.title)
        if case == 'table':
            caption,tab = element
            try:
                print(walkTrace,tab._leopardref,caption)
            except AttributeError:
                tab._leopardref = next(self._reportSection._tabnr)
                print(walkTrace,tab._leopardref,caption)

    def sectionOutZip(self,zipcontainer,zipdir='',figtype='png'):
        """Prepares section for zip output
        """
        from io import StringIO, BytesIO
        text = self.p if not self.settings['doubleslashnewline'] else self.p.replace('//','\n')
        zipcontainer.writestr(
            zipdir+'section.txt',
            '# {}\n{}'.format(self.title,text).encode()
        )
        c = count(1)
        for ftitle,f in self.figs.items():
            figfile = zipdir+'fig{}_{}.{}'.format(next(c),ftitle.replace(' ','_'),figtype)
            b = BytesIO()
            f.savefig(b,format=figtype,transparent=True)
            b.seek(0)
            zipcontainer.writestr(figfile,b.getvalue())
        c = count(1)
        for ttitle,t in self.tabs.items():
            b = StringIO()
            t.to_csv(b,sep=csvsep,decimal=csvdec)
            b.seek(0)
            zipcontainer.writestr(
                zipdir+'table{}_{}.csv'.format(next(c),ttitle.replace(' ','_')),
                b.read().encode()
            )
        c = count(1)
        for s in self.subs:
            s.sectionOutZip(zipcontainer,'{}s{}_{}/'.format(zipdir,next(c),s.title.replace(' ','_')),figtype=figtype)

    @walkerWrapper
    def sectionsPDF(self,walkTrace=tuple(),case=None,element=None,doc=None):
        """Prepares section for PDF output.
        """
        import pylatex as pl
        from .extensions.latex import Verbatim
        if case == 'sectionmain':
            if self.settings['clearpage']: doc.append(pl.utils.NoEscape(r'\clearpage'))
            with doc.create(pl.Section(self.title) if len(walkTrace) == 1 else
                            pl.Subsection(self.title) if len(walkTrace) == 2 else
                            pl.Subsubsection(self.title)):
                text = (self.p.replace('\n',' ').replace('//','\n')
                     if self.settings['doubleslashnewline'] else
                     renewliner(self.p))
                if r'\ref' not in text: doc.append(text)
                else:
                    figrefs = re.compile(r'\\ref\{figref\d+\}')
                    #latexcode = re.compile(r'&@\\.+')
                    lastpos = 0
                    for fr in figrefs.finditer(text):
                        doc.append(text[lastpos:fr.start()])
                        doc.append(pl.utils.NoEscape(text[fr.start():fr.end()]))
                        lastpos = fr.end()
                    doc.append(text[lastpos:])
                if self.code: doc.append(Verbatim(self.code))
                
        if case == 'figure':
            width = r'1\textwidth'
            figtitle,fig = element
            #if fig._suptitle: fig.suptitle('Figure {}: {}'.format(fig.number,fig._suptitle.get_text()))
            #figtitle = fig._suptitle.get_text() if fig._suptitle else ''
            #fig.suptitle('')
            with doc.create(pl.Figure(position='htbp')) as plot:
                plt.figure(fig.number)
                plot.add_plot(width=pl.NoEscape(width))
                plot.add_caption(figtitle)
                plot.append(pl.utils.NoEscape(r'\label{figref'+str(fig.number)+r'}'))
            #fig.suptitle(figtitle if figtitle else None)
            
        if case == 'table':
            caption,t = element
            t = pdSeriesToFrame(t) if type(t) == pd.Series else t
            if self.settings['tablehead']:
                t = t.head(self.settings['tablehead'])
            if self.settings['tablecolumns']:
                t = t[self.settings['tablecolumns']]
            with doc.create(pl.Table(position='ht')) as tablenv:
                tablenv.add_caption(caption)
                with doc.create(pl.Tabular('r|'+'l'*len(t.columns))) as table:
                    table.add_hline()
                    table.add_row(('',)+tuple(t.columns))
                    for row in t.to_records():
                        table.add_row(row)
                    table.add_hline(1)
                    #table.add_empty_row()

    @walkerWrapper
    def sectionsWord(self,walkTrace=tuple(),case=None,element=None,doc=None):
        """Prepares section for word output.
        """
        from docx.shared import Inches
        from io import BytesIO
        #p.add_run('italic.').italic = True
                
        if case == 'sectionmain':
            if self.settings['clearpage']: doc.add_page_break()
            
            doc.add_heading(self.title, level = len(walkTrace))
            for p in renewliner(self.p).split('\n'):
                doc.add_paragraph(p)
                
        if case == 'figure':
            bf=BytesIO()
            figtitle,fig = element
            width = fig.get_size_inches()[0]
            width = Inches(width if width < 6 else 6)
            fig.savefig(bf)
            doc.add_picture(bf, width=width)
            doc.add_heading('Figure {}: {}'.format(
                fig._leopardref,
                figtitle),level=6)
            
        if case == 'table':
            caption,t = element
            tableref = t._leopardref
            t = pdSeriesToFrame(t) if type(t) == pd.Series else t
            if self.settings['tablehead']:
                t = t.head(self.settings['tablehead'])
            if self.settings['tablecolumns']:
                t = t[self.settings['tablecolumns']]

            doc.add_heading('Table {}: {}'.format(
                tableref,
                caption),level=6)
            table = doc.add_table(t.shape[0]+1,t.shape[1]+1)
            for tcell,col in zip(table.rows[0].cells[1:],t.columns):
                tcell.text = str(col)
            for trow,rrow in zip(table.rows[1:],t.to_records()):
                for tcell,rcell in zip(trow.cells,rrow):
                    tcell.text = str(rcell)

    @staticmethod
    def sectionFromFunction(function,*args,**kwargs):
        """
        This staticmethod executes the function that is passed with the provided args and kwargs.
        The first line of the function docstring is used as the section title, the comments
        within the function body are parsed and added as the section text.
        The function should return an ordered dict of figures and tables, that are then
        attached to the section.

        Args:
            function (function): The function that generates the section content.

        Returns:
            Section

        >>> # Section title of example function
        ... def exampleFunction(a,b=None):
        ...     'Mock figures and tables included'
        ...     figures = (('fig1',Mock()),('fig2',Mock()))
        ...     tables = (('tab1',Mock()),('tab2',Mock()))
        ...     return figures, tables
        >>> Section.sectionFromFunction(exampleFunction,Mock(),b=Mock())
        <Section @ Section title of example function>
        """
        figures, tables = function(*args,**kwargs)
        title = inspect.getcomments(function)[1:].strip()
        text = inspect.getdoc(function)
        code = inspect.getsource(function)
        return Section(title=title,text=text,figures=figures,tables=tables,code=code)

class Report(Section):
    """
    Contains as main attribute a list of sections.
    Defines methods of outputting the sections.
    outfile should not include a final extension, as
    that is determined by the different output methods.
    """
    def __init__(
            self, title, intro='', conclusion='', outname='', outfile=None,
            author=None, addTime=True, makeDir=False, capture_print=False):
        """Create a report

        Note:
            If you are using a nested outname and the reports directory does not 
            yet exist, makeDir can be set to True.

        Args:
            title (str): Report title.
            intro (str): Introduction.
            conclustion (str): Report concluding statements.
            outname (str): The base file name for the report, can include 
                directories relative to reportDir.
            outfile (str, optional): A full filename (without extension) can 
                also be provided to save report outside of reportDir.
            capture_print (bool): If True overload print function to report print method
        """
        import time
        super().__init__(title=title,text=intro)
        self.sections = self.subs # Report sections can be accessed by both sections and subs attribute
        self.conclusion = conclusion.strip()
        self.outdir = os.path.dirname(outfile) if outfile else os.path.join(reportsDir,outname)
        self.outfile = outfile if outfile else os.path.join(
            self.outdir,
            '{}{}'.format(time.strftime('%Y_%m_%d'), '_'+outname if outname else '')
        )
        self.author = author
        self.addTime = addTime

        #Fig and table ref management hidden variables
        self._fignr = count(1)
        self._tabnr = count(1)

        if capture_print:
            from .utils import print2report
            print2report(self)

    def print(self, *args, **kwargs):
        """Method that can be used instead of print, to
        attach the stdout to the report as well in the
        lastSection's code attribute.

        This does turn the code attribute to a str, so
        the same section cannot contain a code object.
        """
        from io import StringIO
        import builtins
        print = builtins._print if '_print' in vars(builtins) else builtins.print
        if not 'file' in kwargs:
            out = StringIO()
            print(*args, file=out, **kwargs)
            print(out.getvalue(), end='')
            try: self.lastSection.code += out.getvalue()
            except TypeError: self.lastSection.code = out.getvalue()
            except AttributeError: print(self, 'has no sections yet')
        else: print(*args, **kwargs)
    
    def list(self):
        """
        Get an overview of the report content list
        """
        for i in range(len(self.sections)):
            self.sections[i].list(walkTrace=(i+1,))
        
    def outputZip(self,figtype='png'):
        """
        Outputs the report in a zip container.
        Figs and tabs as pngs and excells.

        Args:
            figtype (str): Figure type of images in the zip folder.
        """
        from zipfile import ZipFile
        with ZipFile(self.outfile+'.zip', 'w') as zipcontainer:
            zipcontainer.writestr(
                'summary.txt',
                '# {}\n\n{}\n{}'.format(
                    self.title,
                    self.p,
                    ('\n## Conclusion\n' if self.conclusion else '')+self.conclusion
                ).encode()
            )
            c = count(1)
            for section in self.sections:
                section.sectionOutZip(zipcontainer,'s{}_{}/'.format(next(c),section.title.replace(' ','_')),
                                      figtype=figtype)

    def outputPDF(self, show = False, geometry_options = None, **kwargs):
        """Makes a pdf report with pylatex
        *kwargs* are send to doc.generate_pdf
        -> see pylatex.Document.generate_pdf for help
        """
        import pylatex as pl
        # geometry_options giving an error on Mac
        if geometry_options and isinstance(geometry_options, bool):
            geometry_options = {"tmargin": "2cm", "lmargin": "2cm"}
        
        doc = pl.Document(geometry_options=geometry_options)
        
        #Following option avoids float error when to many unplaced figs or tabs
        # (to force placing floats also \clearpage can be used after a section for example)
        doc.append(pl.utils.NoEscape(r'\extrafloats{100}'))
        doc.append(pl.utils.NoEscape(r'\title{'+self.title+'}'))
        if self.addTime:
            from time import localtime, strftime
            doc.append(pl.utils.NoEscape(r'\date{'+strftime("%Y-%m-%d %H:%M:%S", localtime())+r'}'))
        else: doc.append(pl.utils.NoEscape(r'\date{\today}'))
        if self.author: doc.append(pl.utils.NoEscape(r'\author{'+self.author+'}'))
        doc.append(pl.utils.NoEscape(r'\maketitle'))

        # Append introduction
        if self.p:
            with doc.create(pl.Section('Introduction')):
                doc.append(
                    self.p.replace('\n',' ').replace('//','\n')
                    if self.settings['doubleslashnewline'] else
                    renewliner(self.p)
                )

        # Sections
        c = count(1)
        for section in self.sections:
            section.sectionsPDF(walkTrace=(next(c),),doc=doc)

        # Append conclusion
        if self.conclusion:
            with doc.create(pl.Section('Conclusion')):
                doc.append(
                    self.conclusion.replace('\n',' ').replace('//','\n')
                    if self.settings['doubleslashnewline'] else
                    renewliner(self.conclusion)
                )

        # Generate pdf
        doc.generate_pdf(self.outfile,**kwargs)

        # Open file
        if show:
            from .utils import open_file
            open_file(self.outfile+'.pdf')

    def outputWord(self):
        """Output report to word docx
        """
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = docx.Document()
        doc.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_heading(self.title, level=0)
        if self.addTime:
            from time import localtime, strftime
            doc.add_heading(strftime("%Y-%m-%d %H:%M:%S", localtime()), level=1)
            
         # Append introduction
        if self.p:
            doc.add_heading('Introduction',level=1)
            for p in renewliner(self.p).split('\n'):
                doc.add_paragraph(p)

        # Sections
        c = count(1)
        #Prepare fig and table numbers
        self.listFigures(tuple())
        self.listTables(tuple())
        for section in self.sections:
            section.sectionsWord((next(c),),doc=doc)

        # Append conclusion
        if self.conclusion:
            doc.add_heading('Conclusion', level=1)
            for p in renewliner(self.conclusion).split('\n'):
                doc.add_paragraph(p)

        # Generate Word document
        doc.save(self.outfile+'.docx')

    @staticmethod
    def getReportTable(reportzipfile,tablefilename,inReportsDir=True,verbose=False):
        """Get a pandas table from a previous report

        Args:
            reportzipfile (str): Zip folder location, '.zip' extension is optional.
            tablefilename (str or list int): Table location within the zip folder.
                Can be provided as the filename within the zip folder, or a list of integers
                indicating its exact position (1-indexed). If you provide an empty string or
                list, all available table filenames in the zip folder will be printed.
            inReportsDir (bool): Search reportzipfile relative to reportsDir.

        Returns:
            pd.DataFrame
        """
        import zipfile, io, re

        # zipfilename preparation
        if not reportzipfile.endswith('.zip'): reportzipfile+='.zip'
        if inReportsDir: reportzipfile = os.path.join(reportsDir,reportzipfile)
        with zipfile.ZipFile(reportzipfile) as z:
            # print all table filenames if tablefilename is not provided
            if not tablefilename:
                for f in z.filelist:
                    if 'table' in f.filename: print(f.filename)
                return
            # tablefilename preparation if int list
            if isinstance(tablefilename,list):
                tablelocation = tablefilename
                tablefilename = None
                location = re.compile(r'(s|table)(\d+)_')
                for f in z.filelist:
                    if 'table' not in f.filename or f.filename.count('/') != (len(tablelocation)-1): continue
                    if [int(location.match(s).groups()[1]) for s in f.filename.split('/')] == tablelocation:
                        tablefilename = f.filename
                        if verbose: print('Loading',tablefilename)
                        break
                if tablefilename is None: raise FileNotFoundError('Table location not found in zip folder.')
            # read table
            with z.open(tablefilename) as f:
                ft = io.TextIOWrapper(f)
                return pd.read_csv(ft,index_col=0,sep=csvsep,decimal=csvdec)

class Presentation(Report):
    """
    Outputs a presentation pdf.
    Currently inherits from Report, so any kind of Sections
    can be appended to it, although only sections should be
    appended that have at most 1 figure or 1 table.
    """
    def outputPDF(self, theme=None, colortheme=None, colorel=None, show=False, **kwargs):
        """Makes a pdf presentation with pylatex
        *kwargs* are send to doc.generate_pdf 
        -> see pylatex.Document.generate_pdf for help
        -> see https://deic-web.uab.cat/~iblanes/beamer_gallery/index.html for (color)theme options

        Args:
          theme (str): beamer theme
          colortheme (str): beamer color theme
          colorel (str | list): color(settings) to color additional elements
            when provided as list, should be of format
            (rgb-tuple, rgb-tuple, [textcolor1-string, [textcolor2-string]])
          show (bool): Display the file generated
        """
        import pylatex as pl
        from .extensions.latex import Frame
        
        #geometry_options = {"tmargin": "2cm", "lmargin": "2cm"}
        doc = pl.Document(
            documentclass='beamer'#, geometry_options=geometry_options
        )
        doc.content_separator = '\n' # can give issues for combi verbatim frame
        if theme: doc.preamble.append(pl.NoEscape(r'\usetheme{'+pl.NoEscape(theme)+pl.NoEscape('}')))
        if colortheme:
            doc.preamble.append(
                pl.NoEscape(r'\usecolortheme{'+pl.NoEscape(colortheme)+pl.NoEscape('}'))
            )
        doc.preamble.append(pl.NoEscape(r'\setbeamertemplate{caption}{\insertcaption}'))
        if colorel and isinstance(colorel, str):
            colorel = colorel if '!' in colorel else f'{colorel}!80!black'
            for element in ('caption', 'itemize item'):
                doc.preamble.append(
                    pl.NoEscape(r'\setbeamercolor{')+
                    pl.NoEscape(element)+
                    pl.NoEscape(r'}{fg=')+
                    pl.NoEscape(colorel)+
                    pl.NoEscape(r'}')
                )
        elif colorel:
            from .extensions.latex import get_beamer_color_preamble
            doc.preamble += get_beamer_color_preamble(*colorel)
        
        doc.append(pl.NoEscape(r'\title{'+self.title+'}'))
        if self.addTime:
            from time import localtime, strftime
            doc.append(pl.utils.NoEscape(r'\date{'+strftime("%Y-%m-%d %H:%M:%S", localtime())+r'}'))
        else: doc.append(pl.utils.NoEscape(r'\date{\today}'))
        if self.author: doc.append(pl.utils.NoEscape(r'\author{'+self.author+'}'))
        #doc.append(pl.utils.NoEscape(r'\maketitle'))
        doc.append(pl.utils.NoEscape(r'\begin{frame} \titlepage \end{frame}'))

        # Sections
        c = count(1)
        for section in self.sections:
            #section.sectionsPDF(walkTrace=(next(c),),doc=doc)
            ncols = len(section.figs)+len(section.tabs)
            frame = Frame(section.title, section.subtitle, ncols=ncols if ncols>1 else 0)
            
            for i,fig in enumerate(section.figs):
                coli = i+len(section.tabs) #TODO add option to use first column(s) for figures
                width = pl.NoEscape(r'1\textwidth')
                figtitle,fig = list(section.figs.items())[i]
                plot = pl.Figure()
                plt.figure(fig.number)
                plot.add_plot(width=width)
                plot.add_caption(figtitle)
                plot.append(pl.utils.NoEscape(r'\label{figref'+str(fig.number)+r'}'))
                if ncols <=1:
                    frame.append(plot)
                else: frame.columns.cols[coli].append(plot)
            
            for i,tab in enumerate(section.tabs):
                #i += len(section.figs) #TODO add option to use first column(s) for tables
                caption,t = list(section.tabs.items())[i]
                t = pdSeriesToFrame(t) if type(t) == pd.Series else t
                tablenv = pl.Table()
                tablenv.add_caption(caption)
                table = pl.Tabular('r|'+'l'*len(t.columns))
                table.add_hline()
                table.add_row(('',)+tuple(t.columns))
                for row in t.to_records():
                    table.add_row(row)
                table.add_hline(1)
                tablenv.append(table)
                if ncols <=1:
                    frame.append(tablenv)
                else: frame.columns.cols[i].append(tablenv)

            if section.p:
                if isinstance(section.p, str):
                    frame.append(section.p)
                elif isinstance(section.p, list):
                    frame.add_enumeration(section.p)

            if section.code:
                if isinstance(section.code, str):
                    frame.add_code(section.code)
                else:
                    import inspect
                    frame.add_code(
                        inspect.getsource(section.code),
                        language='python'
                    )
            
            doc.append(frame)

        # Generate pdf
        doc.generate_pdf(self.outfile,**kwargs)

        # Open file
        if show:
            from .utils import open_file
            open_file(self.outfile+'.pdf')
